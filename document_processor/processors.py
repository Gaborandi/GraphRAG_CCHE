# document_processor/processors.py
from pathlib import Path
import PyPDF2
from docx import Document as DocxDocument
import pandas as pd
import json
from typing import Dict, Any, Union, List
import logging
from datetime import datetime
import hashlib

from .base import Document, DocumentProcessor
from ..config import Config
from ..utils.text_processing import TextPreprocessor

class PDFProcessor(DocumentProcessor):
    """Processor for PDF documents."""
    
    def __init__(self, config: Config):
        super().__init__(config)
        self.preprocessor = TextPreprocessor()

    def process(self, file_path: Union[str, Path]) -> Document:
        """Process a PDF file and return a Document instance."""
        file_path = Path(file_path)
        self.logger.info(f"Processing PDF file: {file_path}")
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract text from all pages
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                
                # Clean extracted text
                cleaned_text = self.preprocessor.clean_text(text)
                
                # Extract metadata
                metadata = self._extract_pdf_metadata(pdf_reader, file_path)
                
                # Create document instance
                doc = Document(
                    content=cleaned_text,
                    metadata=metadata,
                    doc_id=self._generate_doc_id(file_path, cleaned_text),
                    source_path=file_path
                )
                
                # Create chunks
                doc.chunks = self.chunk_text(cleaned_text)
                
                return doc
                
        except Exception as e:
            self.logger.error(f"Error processing PDF {file_path}: {str(e)}")
            raise

    def _extract_pdf_metadata(self, pdf_reader: PyPDF2.PdfReader, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from PDF file."""
        try:
            metadata = {
                'file_type': 'pdf',
                'file_size': file_path.stat().st_size,
                'file_name': file_path.name,
                'creation_date': datetime.fromtimestamp(file_path.stat().st_ctime).isoformat(),
                'modification_date': datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
                'num_pages': len(pdf_reader.pages)
            }
            
            # Add PDF-specific metadata if available
            if pdf_reader.metadata:
                metadata.update({
                    'title': pdf_reader.metadata.get('/Title', ''),
                    'author': pdf_reader.metadata.get('/Author', ''),
                    'subject': pdf_reader.metadata.get('/Subject', ''),
                    'creator': pdf_reader.metadata.get('/Creator', '')
                })
            
            return metadata
            
        except Exception as e:
            self.logger.warning(f"Error extracting PDF metadata: {str(e)}")
            return {'file_type': 'pdf', 'error': str(e)}

class DocxProcessor(DocumentProcessor):
    """Processor for DOCX documents."""
    
    def __init__(self, config: Config):
        super().__init__(config)
        self.preprocessor = TextPreprocessor()

    def process(self, file_path: Union[str, Path]) -> Document:
        """Process a DOCX file and return a Document instance."""
        file_path = Path(file_path)
        self.logger.info(f"Processing DOCX file: {file_path}")
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Clean extracted text
            cleaned_text = self.preprocessor.clean_text(text)
            
            # Extract metadata
            metadata = self._extract_docx_metadata(doc, file_path)
            
            # Create document instance
            document = Document(
                content=cleaned_text,
                metadata=metadata,
                doc_id=self._generate_doc_id(file_path, cleaned_text),
                source_path=file_path
            )
            
            # Create chunks
            document.chunks = self.chunk_text(cleaned_text)
            
            return document
            
        except Exception as e:
            self.logger.error(f"Error processing DOCX {file_path}: {str(e)}")
            raise

    def _extract_docx_metadata(self, doc: DocxDocument, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from DOCX file."""
        try:
            core_properties = doc.core_properties
            metadata = {
                'file_type': 'docx',
                'file_size': file_path.stat().st_size,
                'file_name': file_path.name,
                'creation_date': datetime.fromtimestamp(file_path.stat().st_ctime).isoformat(),
                'modification_date': datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
                'author': core_properties.author or '',
                'title': core_properties.title or '',
                'subject': core_properties.subject or '',
                'revision': core_properties.revision or ''
            }
            return metadata
            
        except Exception as e:
            self.logger.warning(f"Error extracting DOCX metadata: {str(e)}")
            return {'file_type': 'docx', 'error': str(e)}

class CSVProcessor(DocumentProcessor):
    """Processor for CSV documents."""
    
    def __init__(self, config: Config):
        super().__init__(config)
        self.preprocessor = TextPreprocessor()

    def process(self, file_path: Union[str, Path]) -> Document:
        """Process a CSV file and return a Document instance."""
        file_path = Path(file_path)
        self.logger.info(f"Processing CSV file: {file_path}")
        
        try:
            # Read CSV file
            df = pd.read_csv(file_path)
            
            # Convert DataFrame to text representation
            text = self._dataframe_to_text(df)
            
            # Clean text
            cleaned_text = self.preprocessor.clean_text(text)
            
            # Extract metadata
            metadata = self._extract_csv_metadata(df, file_path)
            
            # Create document instance
            document = Document(
                content=cleaned_text,
                metadata=metadata,
                doc_id=self._generate_doc_id(file_path, cleaned_text),
                source_path=file_path
            )
            
            # Create chunks
            document.chunks = self.chunk_text(cleaned_text)
            
            return document
            
        except Exception as e:
            self.logger.error(f"Error processing CSV {file_path}: {str(e)}")
            raise

    def _dataframe_to_text(self, df: pd.DataFrame) -> str:
        """Convert DataFrame to text representation."""
        # Convert DataFrame to JSON-like format
        text_parts = []
        
        # Add header description
        columns = df.columns.tolist()
        text_parts.append(f"CSV file with columns: {', '.join(columns)}")
        
        # Add row descriptions
        for idx, row in df.iterrows():
            row_text = []
            for col in columns:
                value = row[col]
                if pd.notna(value):  # Skip NaN values
                    row_text.append(f"{col}: {value}")
            text_parts.append(" | ".join(row_text))
        
        return "\n".join(text_parts)

    def _extract_csv_metadata(self, df: pd.DataFrame, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from CSV file and DataFrame."""
        try:
            metadata = {
                'file_type': 'csv',
                'file_size': file_path.stat().st_size,
                'file_name': file_path.name,
                'creation_date': datetime.fromtimestamp(file_path.stat().st_ctime).isoformat(),
                'modification_date': datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
                'num_rows': len(df),
                'num_columns': len(df.columns),
                'columns': df.columns.tolist(),
                'data_types': {col: str(dtype) for col, dtype in df.dtypes.items()}
            }
            return metadata
            
        except Exception as e:
            self.logger.warning(f"Error extracting CSV metadata: {str(e)}")
            return {'file_type': 'csv', 'error': str(e)}
